"""
DOCX Resume Parser - Extract text from Word document resumes
"""

from docx import Document
from io import BytesIO
import logging

logger = logging.getLogger(__name__)


def parse_docx_resume(file_bytes: bytes) -> str:
    """
    Extract text from a DOCX resume.
    
    Returns extracted text string.
    """
    try:
        doc = Document(BytesIO(file_bytes))
        
        text_parts = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        
        return "\n".join(text_parts)
    
    except Exception as e:
        logger.error(f"DOCX parsing error: {e}")
        return ""


def extract_docx_structure(file_bytes: bytes) -> dict:
    """Extract structured content from DOCX."""
    try:
        doc = Document(BytesIO(file_bytes))
        
        sections = {}
        current_section = "General"
        sections[current_section] = []
        
        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            
            # Detect headings
            if para.style.name.startswith("Heading"):
                current_section = para.text.strip()
                sections[current_section] = []
            else:
                sections[current_section].append(para.text.strip())
        
        return sections
    except Exception as e:
        logger.error(f"DOCX structure extraction error: {e}")
        return {}